#!/usr/bin/env python3

import docx
from docx import Document
import sys
import os
import re

# Add the current directory to path to import from faq_processor
sys.path.insert(0, os.getcwd())
from faq_processor import extract_paragraph_content

def find_hyperlinks_in_document():
    print("Finding all hyperlinks in the document...")
    
    # Load the data engineering document
    doc = Document('cache/data-engineering-zoomcamp.docx')
    
    hyperlink_count = 0
    paragraph_count = 0
    
    for i, paragraph in enumerate(doc.paragraphs):
        paragraph_has_hyperlink = False
        paragraph_text = paragraph.text
        
        # Skip empty paragraphs
        if not paragraph_text.strip():
            continue
            
        paragraph_count += 1
        
        # Check each run for hyperlinks
        for j, run in enumerate(paragraph.runs):
            element = run._element
            
            # Walk up the element tree to find hyperlinks
            current = element
            while current is not None:
                if current.tag.endswith('}hyperlink'):
                    if not paragraph_has_hyperlink:  # Only print once per paragraph
                        hyperlink_count += 1
                        print(f"\n=== Hyperlink found in Paragraph {i+1} ===")
                        print(f"Style: {paragraph.style.name}")
                        print(f"Full paragraph text: '{paragraph_text}'")
                        
                        # Extract content using our function
                        extracted_content = extract_paragraph_content(paragraph)
                        print(f"Extracted content: '{extracted_content}'")
                        
                        paragraph_has_hyperlink = True
                    
                    # Get hyperlink details
                    rel_id = current.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    print(f"  Run {j} text: '{run.text}'")
                    print(f"  Relationship ID: {rel_id}")
                    
                    try:
                        rel = paragraph._parent.part.rels[rel_id]
                        print(f"  Target URL: {rel.target_ref}")
                    except (KeyError, AttributeError) as e:
                        print(f"  Error getting URL: {e}")
                    
                    break
                current = current.getparent()
        
        # Limit output to first 5 hyperlinks for debugging
        if hyperlink_count >= 5:
            break
    
    print(f"\nSummary: Found {hyperlink_count} hyperlinks in {paragraph_count} non-empty paragraphs")

def find_url_patterns():
    print("\n" + "="*50)
    print("Searching for URL patterns that might not be hyperlinks...")
    
    # Load the data engineering document
    doc = Document('cache/data-engineering-zoomcamp.docx')
    
    url_pattern = re.compile(r'https?://[^\s\]]+')
    url_count = 0
    
    for i, paragraph in enumerate(doc.paragraphs):
        paragraph_text = paragraph.text
        
        # Find URLs in the text
        urls = url_pattern.findall(paragraph_text)
        if urls:
            url_count += len(urls)
            print(f"\n=== URLs found in Paragraph {i+1} ===")
            print(f"Style: {paragraph.style.name}")
            print(f"Text: '{paragraph_text}'")
            print(f"URLs: {urls}")
            
            # Check if any run is marked as hyperlink
            has_hyperlink_markup = False
            for run in paragraph.runs:
                element = run._element
                current = element
                while current is not None:
                    if current.tag.endswith('}hyperlink'):
                        has_hyperlink_markup = True
                        break
                    current = current.getparent()
            
            print(f"Has hyperlink markup: {has_hyperlink_markup}")
            
            # Limit output
            if url_count >= 10:
                break
    
    print(f"\nSummary: Found {url_count} URL patterns")

if __name__ == "__main__":
    find_hyperlinks_in_document()
    find_url_patterns()