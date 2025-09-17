#!/usr/bin/env python3

import docx
from docx import Document
import sys
import os

# Add the current directory to path to import from faq_processor
sys.path.insert(0, os.getcwd())
from faq_processor import extract_paragraph_content

def test_link_extraction():
    print("Testing link extraction...")
    
    # Load the test document
    doc = Document('test_links.docx')
    
    print(f"Document loaded with {len(doc.paragraphs)} paragraphs\n")
    
    for i, paragraph in enumerate(doc.paragraphs):
        print(f"=== Paragraph {i+1} ===")
        print(f"Style: {paragraph.style.name}")
        print(f"Raw text: '{paragraph.text}'")
        
        # Extract content using our function
        extracted_content = extract_paragraph_content(paragraph)
        print(f"Extracted content: '{extracted_content}'")
        
        # Debug: Check runs and their properties
        print(f"Number of runs: {len(paragraph.runs)}")
        for j, run in enumerate(paragraph.runs):
            print(f"  Run {j}: '{run.text}'")
            
            # Check for hyperlink properties
            element = run._element
            parent = element.getparent()
            print(f"    Element tag: {element.tag}")
            print(f"    Parent tag: {parent.tag if parent is not None else 'None'}")
            
            # Look for hyperlink in the element tree
            hyperlink_found = False
            current = element
            while current is not None:
                if current.tag.endswith('}hyperlink'):
                    print(f"    *** HYPERLINK FOUND ***")
                    rel_id = current.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    print(f"    Relationship ID: {rel_id}")
                    hyperlink_found = True
                    
                    # Try to get the actual URL
                    try:
                        rel = paragraph._parent.part.rels[rel_id]
                        print(f"    Target URL: {rel.target_ref}")
                    except (KeyError, AttributeError) as e:
                        print(f"    Error getting URL: {e}")
                    break
                current = current.getparent()
            
            if not hyperlink_found:
                print(f"    No hyperlink found")
        
        print()

if __name__ == "__main__":
    test_link_extraction()