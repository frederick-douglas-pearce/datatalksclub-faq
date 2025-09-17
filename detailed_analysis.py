#!/usr/bin/env python3

import docx
from docx import Document
import sys
import os

def detailed_hyperlink_analysis():
    print("Detailed analysis of test_links.docx...")
    
    try:
        # Load your test document
        doc = Document('test_links.docx')
        
        print(f"Document loaded with {len(doc.paragraphs)} paragraphs\n")
        
        # Also check document relationships
        print("=== Document Relationships ===")
        for rel_id, rel in doc.part.rels.items():
            print(f"Relationship {rel_id}: {rel.target_ref}")
        
        print("\n=== Paragraph Analysis ===")
        
        for i, paragraph in enumerate(doc.paragraphs):
            if not paragraph.text.strip():
                continue
                
            print(f"\n--- Paragraph {i+1} ---")
            print(f"Raw text: '{paragraph.text}'")
            print(f"Number of runs: {len(paragraph.runs)}")
            
            for j, run in enumerate(paragraph.runs):
                print(f"\n  Run {j}:")
                print(f"    Text: '{run.text}'")
                print(f"    Element tag: {run._element.tag}")
                
                # Get all ancestors
                element = run._element
                ancestors = []
                current = element
                while current is not None:
                    ancestors.append(current.tag.split('}')[-1])
                    current = current.getparent()
                
                print(f"    Ancestor chain: {' -> '.join(ancestors)}")
                
                # Check for hyperlink properties in the element tree
                current = element
                hyperlink_info = None
                while current is not None:
                    if current.tag.endswith('}hyperlink'):
                        rel_id = current.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        hyperlink_info = {
                            'rel_id': rel_id,
                            'tag': current.tag
                        }
                        
                        # Try to get actual URL
                        try:
                            rel = paragraph._parent.part.rels[rel_id]
                            hyperlink_info['url'] = rel.target_ref
                        except (KeyError, AttributeError) as e:
                            hyperlink_info['error'] = str(e)
                        break
                    current = current.getparent()
                
                if hyperlink_info:
                    print(f"    HYPERLINK FOUND: {hyperlink_info}")
                else:
                    print(f"    No hyperlink found")
    
    except FileNotFoundError:
        print("test_links.docx not found. Please make sure the file exists.")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    detailed_hyperlink_analysis()