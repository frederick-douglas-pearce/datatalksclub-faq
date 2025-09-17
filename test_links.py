#!/usr/bin/env python3

import docx
from docx import Document
from docx.shared import Inches

# Create a test document with different types of links
doc = Document()

# Add a title
doc.add_heading('Test Document for Link Extraction', 0)

# Add a section heading
doc.add_heading('General Section', level=1)

# Add a question heading  
doc.add_heading('How to test links?', level=2)

# Add paragraphs with different link types
p1 = doc.add_paragraph('This is a plain text reference: Check this article for details - Setting up docker in macOS')

p2 = doc.add_paragraph('This is a hyperlink: ')
hyperlink = p2.add_run('GitHub Issue')
# Add hyperlink properties (this simulates what Word does)

p3 = doc.add_paragraph('This is a URL: https://github.com/Homebrew/brew/issues/16309')

p4 = doc.add_paragraph('Mixed content: See the ')
run = p4.add_run('documentation link')
p4.add_run(' for more details.')

# Save the test document
doc.save('test_links.docx')
print("Created test_links.docx")