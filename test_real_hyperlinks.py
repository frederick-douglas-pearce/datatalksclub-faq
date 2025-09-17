#!/usr/bin/env python3

import docx
from docx import Document
import sys
import os

# Add the current directory to path to import from faq_processor
sys.path.insert(0, os.getcwd())
from faq_processor import extract_paragraph_content

def test_real_hyperlinks():
    print("Testing extraction from your test_links.docx...")
    
    try:
        # Load your test document
        doc = Document('test_links.docx')
        
        print(f"Document loaded with {len(doc.paragraphs)} paragraphs\n")
        
        for i, paragraph in enumerate(doc.paragraphs):
            if not paragraph.text.strip():
                continue
                
            print(f"=== Paragraph {i+1} ===")
            print(f"Style: {paragraph.style.name}")
            print(f"Raw text: '{paragraph.text}'")
            
            # Extract content using our function
            extracted_content = extract_paragraph_content(paragraph)
            print(f"Extracted content: '{extracted_content}'")
            
            # Check for hyperlinks in runs
            hyperlink_found = False
            for j, run in enumerate(paragraph.runs):
                element = run._element
                
                # Walk up the element tree to find hyperlinks
                current = element
                while current is not None:
                    if current.tag.endswith('}hyperlink'):
                        if not hyperlink_found:
                            print(f"*** HYPERLINK DETECTED ***")
                            hyperlink_found = True
                        
                        rel_id = current.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        print(f"  Run {j} text: '{run.text}'")
                        print(f"  Relationship ID: {rel_id}")
                        
                        try:
                            rel = paragraph._parent.part.rels[rel_id]
                            print(f"  Target URL: {rel.target_ref}")
                        except (KeyError, AttributeError) as e:
                            print(f"  Error getting URL: {e}")
                        break
                    current = current.getparent()
            
            if not hyperlink_found:
                print("No hyperlinks found in this paragraph")
            
            print()
    
    except FileNotFoundError:
        print("test_links.docx not found. Please make sure the file exists.")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    test_real_hyperlinks()